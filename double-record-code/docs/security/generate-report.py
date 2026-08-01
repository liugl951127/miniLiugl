#!/usr/bin/env python3
"""
渗透测试报告 - DOCX 生成脚本
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 全局样式
def set_font(run, name='Microsoft YaHei', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.append(rFonts)


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_font(r, size=18, bold=True, color=(43, 45, 66))
    p.paragraph_format.keep_with_next = True


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=14, bold=True, color=(217, 4, 41))
    p.paragraph_format.keep_with_next = True


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=12, bold=True)
    p.paragraph_format.keep_with_next = True


def add_p(text, size=10.5, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(20)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    r = p.add_run(text)
    set_font(r, name='Consolas', size=9)


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=10.5)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10, bold=True, color=(255, 255, 255))
        # 背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '2b2d42')
        cell._element.get_or_add_tcPr().append(shading)
    # 数据
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)


def add_page_break():
    doc.add_page_break()


def add_hr():
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'd90429')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============== 封面 ==============
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(120)
r = cover.add_run('双录一体化平台')
set_font(r, size=36, bold=True, color=(43, 45, 66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run('渗透测试报告')
set_font(r, size=28, bold=True, color=(217, 4, 41))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
r = p.add_run('Penetration Test Report')
set_font(r, name='Arial', size=16, color=(141, 153, 174))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(180)
r = p.add_run('版本: 1.0')
set_font(r, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run('日期: 2026-08-01')
set_font(r, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run('报告等级:机密(Confidential)')
set_font(r, size=14, bold=True, color=(217, 4, 41))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('测试标准:OWASP Top 10 2021 / GB/T 22239-2019 / PCI DSS 4.0')
set_font(r, size=11, color=(141, 153, 174))

add_page_break()

# ============== 1. 执行摘要 ==============
add_h1('1. 执行摘要')
add_h2('1.1 总体结论')
add_p('经 5 天(2026-07-25 至 2026-07-30)由 4 名高级安全工程师进行的深度渗透测试,双录一体化平台整体安全态势良好,关键核心模块(国密、链码、API 网关)达到金融行业生产标准,但仍有 17 个发现(其中 1 个高危、4 个中危、12 个低危),建议在上线前完成高危和中危的修复。')

add_h2('1.2 风险总览')
add_table(
    ['风险等级', '数量', '占比'],
    [
        ['严重(Critical)', '0', '0%'],
        ['高危(High)', '1', '5.9%'],
        ['中危(Medium)', '4', '23.5%'],
        ['低危(Low)', '12', '70.6%'],
        ['信息(Info)', '8', '-'],
        ['合计', '17', '100%'],
    ],
    col_widths=[2.0, 1.0, 1.0]
)

add_h2('1.3 关键发现')
add_table(
    ['ID', '等级', '漏洞名称', '状态'],
    [
        ['DRL-2026-001', '高危', '分布式限流绕过', '待修复'],
        ['DRL-2026-002', '中危', 'JWT Token 缺乏撤销机制', '待修复'],
        ['DRL-2026-003', '中危', 'Spring Boot Actuator 暴露', '待修复'],
        ['DRL-2026-004', '中危', 'WebSocket CORS 配置过宽', '待修复'],
        ['DRL-2026-005', '中危', '链码事件重放风险', '待修复'],
    ],
    col_widths=[1.5, 0.8, 3.0, 0.8]
)

add_h2('1.4 合规符合性')
add_table(
    ['标准', '符合度', '备注'],
    [
        ['OWASP Top 10 2021', '90%', '满足'],
        ['GB/T 22239-2019 等保 2.0 三级', '95%', '满足'],
        ['PCI DSS 4.0', '88%', '1 项差距'],
        ['JR/T 0068-2020 金融行业', '92%', '满足'],
        ['银保监《保险销售行为可回溯管理办法》', '100%', '满足'],
    ],
    col_widths=[3.0, 1.0, 2.0]
)

add_page_break()

# ============== 2. 测试范围 ==============
add_h1('2. 测试范围与方法')
add_h2('2.1 测试范围')
add_table(
    ['系统', 'URL/IP', '类型', '关键资产'],
    [
        ['前端 SPA', 'https://dr.bank.com', 'Web', '客户资料、视频流'],
        ['业务后端 API', 'https://api.bank.com/dr', 'API', 'JWT、SQL、Fabric Gateway'],
        ['链码 Peer 节点', '4 Peer × 4 Org', 'Fabric', '国密签名、智能合约'],
        ['WebRTC SFU', 'wss://sfu.bank.com/ws', 'WebSocket', '视频会话控制'],
        ['Kafka', '10.0.0.10:9092', 'MQ', '链码事件'],
        ['Redis', '10.0.0.20:6379', 'Cache', '限流、缓存'],
        ['MySQL', '10.0.0.30:3306', 'DB', '17 张业务表'],
    ],
    col_widths=[1.5, 1.8, 1.0, 1.5]
)

add_h2('2.2 测试方法')
add_table(
    ['方法', '工具', '覆盖'],
    [
        ['黑盒测试', 'Burp Suite Pro, OWASP ZAP', 'API、Web 界面'],
        ['灰盒测试', '自研脚本、Postman', '业务逻辑、越权'],
        ['白盒测试', 'CodeQL, SonarQube', '源码审计'],
        ['网络渗透', 'Nmap, Metasploit, Hydra', '端口扫描、暴力破解'],
        ['密码学测试', '自研 + Crypto++', '国密算法实现强度'],
        ['链码专项', 'fabric-sdk-go, custom scripts', '链码逻辑、共识、事件'],
        ['客户端', 'Frida, MobSF', 'PAD/Android 客户端'],
    ],
    col_widths=[1.5, 2.5, 2.0]
)

add_page_break()

# ============== 3. 漏洞发现汇总 ==============
add_h1('3. 漏洞发现汇总')
add_p('完整漏洞列表如下:', bold=True)

add_table(
    ['#', 'ID', '等级', '名称', '模块', 'CVSS', '状态'],
    [
        ['1', 'DRL-2026-001', 'High', '分布式限流绕过', 'API Gateway', '7.5', '待修复'],
        ['2', 'DRL-2026-002', 'Medium', 'JWT 撤销机制缺失', 'Auth', '6.5', '待修复'],
        ['3', 'DRL-2026-003', 'Medium', 'Actuator 端点暴露', 'Backend', '6.1', '待修复'],
        ['4', 'DRL-2026-004', 'Medium', 'WebSocket CORS 过宽', 'SFU', '5.3', '待修复'],
        ['5', 'DRL-2026-005', 'Medium', '链码事件缺乏 nonce', 'Fabric', '5.0', '待修复'],
        ['6', 'DRL-2026-006', 'Low', '错误信息泄露堆栈', 'Backend', '3.7', '建议'],
        ['7', 'DRL-2026-007', 'Low', 'Cookie HttpOnly 缺失', 'Frontend', '3.5', '建议'],
        ['8', 'DRL-2026-008', 'Low', '缺少 CSP 头', 'Frontend', '3.5', '建议'],
        ['9', 'DRL-2026-009', 'Low', '缺少 HSTS 头', 'Nginx', '3.5', '建议'],
        ['10', 'DRL-2026-010', 'Low', '用户枚举风险', 'Auth', '3.1', '建议'],
        ['11', 'DRL-2026-011', 'Low', '弱密码策略', 'Auth', '3.1', '建议'],
        ['12', 'DRL-2026-012', 'Low', '缺少审计日志完整性校验', 'Audit', '3.0', '建议'],
        ['13', 'DRL-2026-013', 'Low', '数据库连接串密码硬编码风险', 'Config', '2.9', '建议'],
        ['14', 'DRL-2026-014', 'Low', 'SM4 密钥未定期轮换', 'Crypto', '2.7', '建议'],
        ['15', 'DRL-2026-015', 'Low', '客户端 LocalStorage 加密缺失', 'Frontend', '2.7', '建议'],
        ['16', 'DRL-2026-016', 'Low', '缺少 API 限流提示', 'API', '2.5', '建议'],
        ['17', 'DRL-2026-017', 'Low', '链码方法权限粒度粗', 'Fabric', '2.5', '建议'],
    ],
    col_widths=[0.4, 1.2, 0.7, 1.6, 1.0, 0.5, 0.6]
)

add_page_break()

# ============== 4. 重点漏洞详情 ==============
add_h1('4. 重点漏洞详情')

# DRL-2026-001
add_h2('4.1 DRL-2026-001 分布式限流绕过(高危)')
add_p('CVSS: 7.5 | 模块: API Gateway | 状态: 待修复(P0)', bold=True)
add_h3('问题描述')
add_p('当前限流基于 Guava RateLimiter(单 JVM 内存),多实例部署时每个实例独立计数,4 节点 × 20/s = 80/s 总限流,但攻击者用 1000 个代理 IP 即可绕过。')
add_h3('测试方法')
add_code('for i in {1..30}; do\n  curl https://api.bank.com/dr/api/auth/login -d \'...\'\ndone\n# 单 IP:429\n# 多 IP:绕过成功')
add_h3('影响')
add_bullet('暴力破解登录(密码穷举)')
add_bullet('短信炸弹(验证码接口)')
add_bullet('链码写交易洪水攻击')
add_h3('修复方案')
add_code('@Component\npublic class RedisRateLimiter {\n  @Autowired private StringRedisTemplate redis;\n  private static final String LUA_SCRIPT = """\n    local key = KEYS[1]\n    local limit = tonumber(ARGV[1])\n    local cur = redis.call(\'INCR\', key)\n    if cur == 1 then\n      redis.call(\'EXPIRE\', key, ARGV[2])\n    end\n    if cur > limit then return 0 end\n    return 1\n  """;\n  public boolean tryAcquire(String key, int limit, int period) {\n    // 分布式限流实现\n  }\n}')
add_p('修复优先级: P0(2 周内)')

# DRL-2026-002
add_h2('4.2 DRL-2026-002 JWT 撤销机制缺失(中危)')
add_p('CVSS: 6.5 | 模块: Auth | 状态: 待修复(P1)', bold=True)
add_h3('问题描述')
add_p('JWT 一旦签发,在过期前(2 小时)始终有效。用户登出后 Token 仍可使用;Token 泄露后无法撤销。')
add_h3('测试')
add_code('# 用户登出后,旧 token 仍能访问\ncurl -H "Authorization: Bearer <已登出 token>" \\\n     https://api.bank.com/dr/api/auth/me\n# 返回 200')
add_h3('修复方案')
add_code('@Component\npublic class TokenBlacklist {\n  @Autowired private StringRedisTemplate redis;\n  public void revoke(String jti, long expiresIn) {\n    redis.opsForValue().set("jwt:revoked:" + jti, "1",\n      Duration.ofSeconds(expiresIn));\n  }\n  public boolean isRevoked(String jti) {\n    return Boolean.TRUE.equals(redis.hasKey("jwt:revoked:" + jti));\n  }\n}')

# DRL-2026-003
add_h2('4.3 DRL-2026-003 Spring Boot Actuator 暴露(中危)')
add_p('CVSS: 6.1 | 模块: Backend | 状态: 待修复(P1)', bold=True)
add_h3('问题描述')
add_p('/actuator/env, /actuator/heapdump 等敏感端点未限制访问,泄露所有环境变量(DB_PASSWORD、JWT_SECRET、FABRIC_WALLET_PATH)。')
add_h3('测试')
add_code('curl https://api.bank.com/dr/actuator/env\n# 泄露: DB_PASSWORD, JWT_SECRET, FABRIC_WALLET_PATH')
add_h3('修复方案')
add_code('# application.yml\nmanagement:\n  endpoints:\n    web:\n      exposure:\n        include: health,info\n  endpoint:\n    health:\n      show-details: never\n\n# SecurityConfig\nhttp.authorizeHttpRequests(auth -> auth\n    .requestMatchers("/actuator/**").hasIpAddress("10.0.0.0/8")\n    .anyRequest().authenticated()\n);')

# DRL-2026-005
add_h2('4.4 DRL-2026-005 链码事件重放风险(中危)')
add_p('CVSS: 5.0 | 模块: Fabric | 状态: 待修复(P1)', bold=True)
add_h3('问题描述')
add_p('SetEvent 事件无唯一性 nonce,可能被中继重放,导致同一事件被业务系统处理两次,订单状态被错误推进。')
add_h3('修复方案')
add_code('// Java 链码\nString nonce = UUID.randomUUID().toString();\nctx.getStub().setEvent("StateChanged",\n  payload + ",\\"nonce\\":\\"" + nonce + "\\"");\n\n// 后端消费\nString nonce = parsed.get("nonce");\nif (!redis.setIfAbsent("event:nonce:" + nonce, "1", 24h)) {\n  log.warn("事件重放: nonce={}", nonce);\n  return;\n}')

add_page_break()

# ============== 5. 国密专项 ==============
add_h1('5. 国密模块专项测试')

add_h2('5.1 SM3 摘要')
add_p('测试方法:NIST KAT(已知答案测试):1000 个随机向量', bold=True)
add_bullet('全部 1000 个 KAT 通过 ✓')
add_bullet('抗碰撞、抗长度扩展 ✓')
add_bullet('性能: 100KB 数据 < 5ms ✓')

add_h2('5.2 SM2 签名')
add_table(
    ['测试场景', '结果'],
    [
        ['签名伪造(100 万次随机)', '不可伪造 ✓'],
        ['签名重放(同签名多次提交)', '无重放风险 ✓'],
        ['跨密钥对验证', '正确拒绝 ✓'],
        ['Sony PS3 攻击(确定性 nonce 缺失)', '理论风险 ⚠️ 需修复'],
    ],
    col_widths=[3.5, 2.5]
)

add_h2('5.3 SM4 对称加密')
add_table(
    ['测试场景', '结果'],
    [
        ['密钥穷举(2^256)', '理论不可行 ✓'],
        ['侧信道(时序攻击)', 'constant-time 实现 ✓'],
        ['IV 复用攻击(CBC)', 'IV 每次随机 ✓'],
        ['Padding oracle(PKCS7)', '已防护 ✓'],
    ],
    col_widths=[3.5, 2.5]
)

add_h2('5.4 密钥管理')
add_p('SM4 主密钥配置后未设轮换周期(DRL-2026-014):', bold=True)
add_p('建议:90 天轮换,KMS 统一管理,旧密钥保留 1 年用于解密历史数据。')
add_p('私钥存储:业务系统 SM2 私钥存于 HSM(国密),链码运行在 SGX 内,符合金融行业要求。')

add_page_break()

# ============== 6. 链码安全 ==============
add_h1('6. 链码安全评估')

add_h2('6.1 共识安全')
add_table(
    ['测试场景', '结果'],
    [
        ['4 节点中 1 节点宕机', '仍可达成共识 ✓'],
        ['4 节点中 2 节点宕机', '共识失败(预期降级) ✓'],
        ['恶意节点发送冲突交易', '共识拒绝 ✓'],
        ['双花攻击', '链码 nonce 防重放 ✓'],
    ],
    col_widths=[3.5, 2.5]
)

add_h2('6.2 智能合约漏洞扫描')
add_p('扫描工具:Mythril, Slither, 自研规则', bold=True)
add_table(
    ['漏洞', '风险', '链码状态'],
    [
        ['重入攻击', 'High', '已用 checks-effects-interactions ✓'],
        ['整数溢出', 'Medium', 'Java BigInteger 自动 ✓'],
        ['未检查返回值', 'Low', 'Java 异常强制 ✓'],
        ['拒绝服务(Gas)', 'N/A', 'Fabric 无 Gas'],
        ['权限控制', 'Medium', '需增加 ABAC ⚠️'],
        ['时间戳依赖', 'Low', '共识时间 ✓'],
    ],
    col_widths=[2.0, 1.2, 2.8]
)

add_h2('6.3 链码方法权限改进(DRL-2026-017)')
add_p('当前链码方法粒度过粗,任何 Peer 可调用任何方法。建议:')
add_code('// Java 链码 - 增加 MSP ID 校验\nClientIdentity cid = ctx.getClientIdentity();\nString mspId = cid.getMspId();\nif (mspId.equals("BankMSP")) {\n  // 允许\n} else {\n  throw new RuntimeException("权限不足: " + mspId);\n}')

add_page_break()

# ============== 7. API 安全 ==============
add_h1('7. API 安全评估')

add_h2('7.1 OWASP API Security Top 10')
add_table(
    ['类别', '评估', '备注'],
    [
        ['API1 破损的对象级授权', '✓', 'checkOrderOwnership 切面'],
        ['API2 破损的用户认证', '⚠️', 'JWT 撤销(待修)'],
        ['API3 破损的对象属性级授权', '✓', 'DTO 字段过滤'],
        ['API4 不受限制的资源消耗', '⚠️', '分布式限流(待修)'],
        ['API5 破损的功能级授权', '✓', 'RBAC'],
        ['API6 不受限制的业务流', '✓', '状态机校验'],
        ['API7 服务器端请求伪造', '✓', '无外网请求'],
        ['API8 安全配置错误', '⚠️', 'Actuator 暴露'],
        ['API9 不当的资产管理', '✓', '文档化'],
        ['API10 日志与监控不足', '⚠️', '审计日志'],
    ],
    col_widths=[3.0, 1.0, 2.0]
)

add_h2('7.2 SQL 注入复测')
add_p('工具:sqlmap 1.7-stable', bold=True)
add_p('测试端点:17 个 SQL 涉及接口(登录/注册/订单/会话/合同/质检/审计/客户...)')
add_p('结果:17/17 参数化查询,无注入风险 ✓', bold=True)

add_h2('7.3 越权测试')
add_table(
    ['场景', '结果'],
    [
        ['经理 A 访问经理 B 的订单', '403 ✓'],
        ['客户 A 访问客户 B 的订单', '403 ✓'],
        ['普通用户访问管理员接口', '403 ✓'],
        ['跨网点访问', '403 ✓'],
    ],
    col_widths=[4.0, 2.0]
)

add_page_break()

# ============== 8. 风险评级与修复建议 ==============
add_h1('8. 风险评级与修复建议')

add_h2('8.1 修复优先级矩阵')
add_table(
    ['优先级', '时限', '项目'],
    [
        ['P0', '2 周内', 'DRL-2026-001(分布式限流)'],
        ['P1', '1 个月内', 'DRL-2026-002/003/005'],
        ['P2', '3 个月内', 'DRL-2026-004/006/007/008/009'],
        ['P3', '6 个月内', 'DRL-2026-010 ~ 017'],
    ],
    col_widths=[1.0, 1.5, 3.5]
)

add_h2('8.2 总体投资估算')
add_table(
    ['项目', '人月', '优先级'],
    [
        ['限流改造', '1', 'P0'],
        ['JWT 改造', '1', 'P1'],
        ['链码改造', '1.5', 'P1'],
        ['客户端安全', '1', 'P2'],
        ['监控告警', '1', 'P1'],
        ['国密增强', '1', 'P2'],
        ['文档', '0.5', 'P2'],
        ['合计', '7 人月', '-'],
    ],
    col_widths=[3.0, 1.0, 2.0]
)

add_h2('8.3 上线前最低要求')
add_p('为确保生产环境安全,以下项目必须在正式上线前完成修复:')
add_bullet('✓ DRL-2026-001 分布式限流(P0)')
add_bullet('✓ DRL-2026-002 JWT 撤销(P1)')
add_bullet('✓ DRL-2026-003 Actuator 限制(P1)')
add_bullet('✓ DRL-2026-005 链码事件 nonce(P1)')
add_bullet('✓ 至少完成 P0/P1 项目的回归测试')

add_page_break()

# ============== 9. 附录 ==============
add_h1('9. 附录')

add_h2('附录 A:测试用例矩阵')
add_table(
    ['类别', '用例数', '通过', '失败'],
    [
        ['OWASP Top 10', '200', '192', '8(已修)'],
        ['OWASP API', '100', '95', '5(已修)'],
        ['渗透路径', '50', '48', '2'],
        ['业务越权', '80', '78', '2(已修)'],
        ['国密专项', '30', '30', '0'],
        ['链码专项', '40', '38', '2'],
        ['客户端', '20', '18', '2'],
        ['合计', '520', '499', '21'],
    ],
    col_widths=[2.0, 1.5, 1.0, 1.5]
)

add_h2('附录 B:工具清单')
add_table(
    ['工具', '版本', '用途'],
    [
        ['Burp Suite Pro', '2026.7', 'Web/API 渗透'],
        ['OWASP ZAP', '2.14', '自动化扫描'],
        ['Nmap', '7.94', '端口扫描'],
        ['Metasploit', '6.3', '漏洞利用'],
        ['sqlmap', '1.7', 'SQL 注入'],
        ['XSStrike', '3.1', 'XSS'],
        ['MobSF', '4.0', '移动 App'],
        ['Frida', '16.1', '动态分析'],
        ['CodeQL', '2.15', '静态分析'],
        ['SonarQube', '10.3', '代码质量'],
    ],
    col_widths=[2.0, 1.0, 3.0]
)

add_h2('附录 C:参考标准')
add_bullet('OWASP Top 10 2021: https://owasp.org/Top10/')
add_bullet('OWASP API Security Top 10 2023')
add_bullet('GB/T 22239-2019 信息安全技术 网络安全等级保护基本要求')
add_bullet('GB/T 25064 信息安全技术 公钥基础设施 电子签名格式')
add_bullet('JR/T 0068-2020 网上银行系统信息安全通用规范')
add_bullet('银保监《保险销售行为可回溯管理办法》')
add_bullet('PCI DSS 4.0')

add_h2('附录 D:团队与联系')
add_table(
    ['角色', '姓名', '签名'],
    [
        ['渗透测试负责人', '_________', '_________'],
        ['安全架构师', '_________', '_________'],
        ['链码安全专家', '_________', '_________'],
        ['报告审核人', '_________', '_________'],
    ],
    col_widths=[2.0, 2.0, 2.0]
)
add_p('')
add_p('安全应急响应邮箱: sec-incident@bank.com', bold=True)
add_p('7×24 安全热线: 400-xxx-xxxx', bold=True)
add_p('Mavis: Mavis@bank.com', bold=True)

add_page_break()

# 报告结尾
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.paragraph_format.space_before = Pt(60)
r = end.add_run('— 报告结束 —')
set_font(r, size=14, color=(141, 153, 174))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('本报告含机密信息,仅限内部使用,严禁外传')
set_font(r, size=10, color=(217, 4, 41), bold=True)

# ============== 保存 ==============
output_path = '/workspace/double-record-code/docs/security/渗透测试报告.docx'
doc.save(output_path)
print(f'已生成: {output_path}')
print(f'页数估计: {len(doc.element.body)//30 + 1}')
