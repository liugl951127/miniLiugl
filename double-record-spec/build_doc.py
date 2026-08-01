"""
Generate the complete technical specification document (.docx) for the
dual-record integration platform.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# ─────────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────────
DIAG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "diagrams")
OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output",
                       "双录一体化平台技术方案说明书.docx")

# Color palette (consistent with diagrams)
PRIMARY   = RGBColor(0x2b, 0x2d, 0x42)
ACCENT    = RGBColor(0xd9, 0x04, 0x29)
SECONDARY = RGBColor(0x8d, 0x99, 0xae)
DARK      = RGBColor(0x1a, 0x1c, 0x2e)
LIGHT     = RGBColor(0xed, 0xf2, 0xf4)
SUCCESS   = RGBColor(0x2A, 0x9D, 0x8F)
WHITE     = RGBColor(0xff, 0xff, 0xff)
GRAY_TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED     = RGBColor(0x6c, 0x75, 0x7d)

# Chinese font
CN_FONT = "Microsoft YaHei"
EN_FONT = "Arial"

# ─────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────
def set_run_font(run, font_size=10.5, bold=False, color=None, font_cn=CN_FONT, font_en=EN_FONT, italic=False):
    run.font.name = font_en
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Set Chinese font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)

def add_paragraph_text(doc, text, size=10.5, bold=False, color=None,
                       align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
                       line_spacing=1.5, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    set_run_font(run, font_size=size, bold=bold, color=color, italic=italic)
    return p

def add_heading(doc, text, level=1):
    """Add styled heading."""
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, font_size=18, bold=True, color=PRIMARY)
        # Bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2b2d42')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, font_size=14, bold=True, color=PRIMARY)
    elif level == 3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, font_size=12, bold=True, color=DARK)
    elif level == 4:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, font_size=11, bold=True, color=ACCENT)

def add_callout(doc, label, text, color=ACCENT):
    """Add a colored callout/note box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add background
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'fce8eb')  # light red bg
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for side in ['left']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '18')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'd90429')
        pBdr.append(b)
    pPr.append(pBdr)
    run1 = p.add_run(f"【{label}】 ")
    set_run_font(run1, font_size=10, bold=True, color=color)
    run2 = p.add_run(text)
    set_run_font(run2, font_size=10, color=GRAY_TEXT)

def add_figure(doc, path, caption, width=6.0):
    """Add image with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    set_run_font(r, font_size=9, italic=True, color=MUTED)

def add_bullet_list(doc, items, indent_level=0):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.left_indent = Pt(20 + 12 * indent_level)
        run = p.add_run(item)
        set_run_font(run, font_size=10.5, color=GRAY_TEXT)

def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.4
        run = p.add_run(item)
        set_run_font(run, font_size=10.5, color=GRAY_TEXT)

def add_table_styled(doc, headers, rows, col_widths=None):
    """Create a styled table with dark blue header and banded rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, font_size=10.5, bold=True, color=WHITE)
        # Cell shading
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2b2d42')
        tcPr.append(shd)
        # Vertical center
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, font_size=10, color=GRAY_TEXT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Banded rows
            if r_idx % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'f5f7fa')
                tcPr.append(shd)

    # Add space after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    return table

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_section_break(doc):
    """Add a new section (next page)."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    return new_section

def setup_default_styles(doc):
    """Set up document defaults: page size, margins, default fonts."""
    section = doc.sections[0]
    # A4 paper
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # 2.5cm margins
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    # Set default style
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)

def add_header_footer(doc, header_text, footer_left):
    section = doc.sections[0]
    # Header
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(header_text)
    set_run_font(run, font_size=9, color=SECONDARY)
    # Footer with page number
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Tab stops for left/center/right
    run = p.add_run(footer_left)
    set_run_font(run, font_size=9, color=SECONDARY)
    run = p.add_run("\t\t")
    # Page number field
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    set_run_font(run, font_size=9, color=SECONDARY)
    run = p.add_run(" / ")
    set_run_font(run, font_size=9, color=SECONDARY)
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"
    run._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar2)
    set_run_font(run, font_size=9, color=SECONDARY)

# ─────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────
def build():
    doc = Document()
    setup_default_styles(doc)
    add_header_footer(doc, "线上线下双录一体化技术方案说明书",
                      "V1.0  |  2026")

    # ══════ Cover Page ══════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("线上线下双录一体化")
    set_run_font(run, font_size=36, bold=True, color=PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("技术方案说明书")
    set_run_font(run, font_size=24, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical Specification for Dual-Record Integration Platform")
    set_run_font(run, font_size=14, italic=True, color=SECONDARY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本:V1.0    |    密级:内部公开    |    2026")
    set_run_font(run, font_size=12, color=GRAY_TEXT)

    # Cover info table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—————————————————————————————")
    set_run_font(run, font_size=10, color=SECONDARY)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_data = [
        ("文档名称",   "线上线下双录一体化技术方案说明书"),
        ("版本号",     "V1.0"),
        ("编写日期",   "2026 年 8 月"),
        ("适用读者",   "技术评审委员会 / 开发团队 / 合规审计 / 业务方"),
    ]
    for i, (k, v) in enumerate(info_data):
        c1 = info_table.rows[i].cells[0]
        c1.text = ""
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(k + "  ")
        set_run_font(run, font_size=10.5, bold=True, color=PRIMARY)
        c1.width = Inches(1.5)

        c2 = info_table.rows[i].cells[1]
        c2.text = ""
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(v)
        set_run_font(run, font_size=10.5, color=GRAY_TEXT)
        c2.width = Inches(4.5)

    add_page_break(doc)

    # ══════ Document Info Page ══════
    add_heading(doc, "文档信息", level=1)

    add_heading(doc, "修订历史", level=2)
    add_table_styled(doc,
        ["版本", "日期", "作者", "修订内容"],
        [
            ["V0.1", "2026-06", "[作者]", "初稿,完成整体架构设计"],
            ["V0.5", "2026-07", "[作者]", "补充话术模板、视频合规、数据一致性方案"],
            ["V0.9", "2026-08", "[作者]", "完成业务流程图、状态机、Saga 补偿机制"],
            ["V1.0", "2026-08", "[作者]", "技术评审通过,正式发布"],
        ],
        col_widths=[0.8, 1.2, 1.2, 3.8])

    add_heading(doc, "术语定义", level=2)
    terms = [
        ["双录", "录音录像", "Insurance Dual-Record;指保险公司、保险中介机构在销售保险产品时,对销售过程同步录音录像的合规要求。"],
        ["话术", "Script", "客户经理销售过程中必须按照监管要求逐字逐句陈述的内容,包括风险揭示、产品告知、客户确认等。"],
        ["适当性管理", "Suitability", "根据客户风险等级、财务状况、投资经验等,推荐匹配的产品。"],
        ["KYC", "Know Your Customer", "了解你的客户,金融机构反洗钱和风险评估的基础制度。"],
        ["Saga", "—", "分布式事务解决方案,通过补偿机制保证最终一致性,适合长事务场景。"],
        ["BPMN", "Business Process Model and Notation", "业务流程模型与标记法,业界标准的流程建模语言。"],
        ["CA 证书", "Certificate Authority", "数字证书认证机构,提供电子签名合规凭证。"],
        ["OCR", "Optical Character Recognition", "光学字符识别,用于证件信息的自动提取。"],
        ["活体检测", "Liveness Detection", "通过动作、表情等判断镜头前是否为真人,防止视频欺诈。"],
        ["区块链存证", "Blockchain Notarization", "利用区块链不可篡改特性固化电子证据,具备司法效力。"],
    ]
    add_table_styled(doc,
        ["术语", "英文/缩写", "说明"],
        terms,
        col_widths=[1.2, 1.4, 4.4])

    add_page_break(doc)

    # ══════ Table of Contents ══════
    add_heading(doc, "目  录", level=1)
    toc_items = [
        ("第 1 章 项目概述", 5),
        ("第 2 章 现状与痛点分析", 6),
        ("第 3 章 整体技术方案", 8),
        ("第 4 章 接入层详细设计", 11),
        ("第 5 章 流程编排层", 12),
        ("第 6 章 能力中台层", 14),
        ("第 7 章 数据治理层", 18),
        ("第 8 章 业务流程详细设计", 20),
        ("第 9 章 话术模板详细设计", 23),
        ("第 10 章 视频合规方案", 26),
        ("第 11 章 智能质检方案", 28),
        ("第 12 章 安全与合规", 30),
        ("第 13 章 接口设计", 32),
        ("第 14 章 实施计划", 34),
        ("附录 A 数据字典", 36),
        ("附录 B 话术模板清单", 37),
        ("附录 C 缩略语表", 38),
    ]
    for t, p in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.tab_stops.add_tab_stop(Inches(6.0),
                                                     alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                                                     leader=2)  # dot leader
        run = para.add_run(t)
        set_run_font(run, font_size=11, color=GRAY_TEXT)
        run = para.add_run("\t" + str(p))
        set_run_font(run, font_size=11, color=PRIMARY, bold=True)

    add_page_break(doc)

    # ══════ Chapter 1: 项目概述 ══════
    add_heading(doc, "第 1 章  项目概述", level=1)

    add_heading(doc, "1.1 项目背景", level=2)
    add_paragraph_text(doc,
        "随着银保监会、证监会等监管机构对金融机构销售合规要求的持续收紧,以及《保险销售行为可回溯管理暂行办法》《证券基金销售机构投资者适当性管理》等法规的深入实施,销售录音录像(以下简称“双录”)已成为银行、保险、证券、基金等行业的强制要求。")
    add_paragraph_text(doc,
        "在传统业务模式下,客户经理在线下高柜或上门办理业务时,通过双录一体机或 PAD 进行本地化录制;线上场景下则通过 H5/小程序内置音视频 SDK 完成录制。两条线长期并行,导致话术版本不一致、视频合规证据链不统一、质检结果无法对账、数据状态分散等系统性痛点。")
    add_paragraph_text(doc,
        "本项目旨在建设统一的“线上线下双录一体化平台”,从架构层面彻底解决上述痛点,实现一套话术、一份视频证据、一套质检规则、一份数据状态贯穿全场景,确保合规零处罚的同时,显著提升客户体验和机构运营效率。")

    add_heading(doc, "1.2 项目目标", level=2)
    add_numbered_list(doc, [
        "统一话术:线上线下共享同一份话术,版本强一致,产品迭代秒级生效;",
        "视频合规:全场景录制加密 → 存储 → 区块链存证 → 监管上报,证据链可追溯;",
        "质检一致:同一规则 + 同一 AI 模型 + 同一评分卡,线上线下结果可比、可对账;",
        "数据原子:分布式事务 + 状态机,任一环节失败自动回滚,不留数据空洞;",
        "客户体验:线上 5-8 分钟完成,线下 15-20 分钟,异常情况自动恢复或人工兜底;",
        "架构先进:5 层分层 + 5 个能力中心,横切能力贯穿,支持未来 3 年业务扩展。",
    ])

    add_heading(doc, "1.3 适用范围", level=2)
    add_paragraph_text(doc,
        "本方案适用于本行(及子公司、关联机构)所有涉及双录合规要求的业务场景,包括但不限于:银行自有理财产品销售、代理保险业务、基金代销业务、信托产品销售、贵金属销售,以及未来可能涉及的证券业务。每类业务对应不同的话术模板、风评要求和监管报送规则,均在本平台统一管理。")

    add_heading(doc, "1.4 参考文档", level=2)
    add_bullet_list(doc, [
        "《保险销售行为可回溯管理暂行办法》(保监发〔2017〕54 号)",
        "《商业银行理财业务监督管理办法》(银保监会 2018 年第 6 号令)",
        "《证券期货投资者适当性管理办法》(证监会令第 130 号)",
        "《金融机构客户身份识别和客户身份资料及交易记录保存管理办法》",
        "本行《信息技术架构白皮书 V3.0》",
        "本行《数据安全管理办法》《个人信息保护管理办法》",
    ])

    add_page_break(doc)

    # ══════ Chapter 2: 现状与痛点分析 ══════
    add_heading(doc, "第 2 章  现状与痛点分析", level=1)

    add_heading(doc, "2.1 线上双录现状", level=2)
    add_paragraph_text(doc,
        "线上双录基于 H5/小程序承载,客户通过手机或 PC 端自助完成全流程。技术栈以 WebRTC 实时音视频 + ASR 语音转写 + 服务端智能质检为主,具备较好的自动化能力,但存在以下问题:")
    add_bullet_list(doc, [
        "话术版本由前端工程团队维护,产品部门更新话术后,线上版本同步周期普遍在 3-7 天;",
        "线上 H5 与线下 PAD/一体机使用不同的话术源,口径不一致曾导致监管检查出现 2 起风险事件;",
        "AI 质检覆盖率达 95% 以上,但质检规则和模型与线下完全独立,两份结果无法直接对比;",
        "流程数据沉淀在线上中台,与线下 CRM/核心系统通过定时同步,实时性差。",
    ])

    add_heading(doc, "2.2 线下双录现状", level=2)
    add_paragraph_text(doc,
        "线下双录由高柜双录一体机或客户经理 PAD 移动展业完成,具备面对面销售的真实性优势,但合规风险更为突出:")
    add_bullet_list(doc, [
        "客户经理需背诵话术,执行偏差大,人工抽检覆盖率不足 5%;",
        "视频文件依赖设备本地存储,U 盘拷贝 + 上传 OSS,介质不一、容易丢失;",
        "纸质单据与电子视频并行,签字、勾选与视频时间戳难以严格对齐;",
        "若视频文件在传输或备份过程中被篡改,事后发现概率极低,合规证据链不完整;",
        "客户必须在网点完成全流程,若中途离席,需重新预约,体验差。",
    ])

    add_heading(doc, "2.3 核心痛点清单", level=2)
    add_callout(doc, "痛点 1", "话术互通:线上/线下两套话术,产品迭代后不同步,口径不一致引发监管风险。")
    add_callout(doc, "痛点 2", "视频合规:线下视频易篡改/丢失/介质不一,合规证据链不完整,事后追溯难。")
    add_callout(doc, "痛点 3", "质检一致:线上 AI 质检、线下人工抽检,两套标准两套结果,无法闭环。")
    add_callout(doc, "痛点 4", "数据原子:双录视频、订单、合同、CRM 状态分散,任一环节失败导致数据不一致。")

    add_heading(doc, "2.4 业务影响分析", level=2)
    add_paragraph_text(doc,
        "上述痛点已在多个维度形成实质影响:监管层面,2025 年内已出现 2 次监管问询;客户层面,中途中断重做率高达 12%,NPS 下降 8 分;机构层面,客户经理人均产能因合规工时被压缩 30%。本方案的设计目标即为系统性解决上述问题。")

    add_page_break(doc)

    # ══════ Chapter 3: 整体技术方案 ══════
    add_heading(doc, "第 3 章  整体技术方案", level=1)

    add_heading(doc, "3.1 设计原则", level=2)
    add_numbered_list(doc, [
        "统一性:一套业务逻辑贯穿所有接入端,避免任何形式的“线上版本”和“线下版本”分裂;",
        "可追溯:全链路埋点 + 区块链存证,任意一笔订单均可完整回溯;",
        "可降级:任一非核心模块故障,均不影响核心业务继续进行;",
        "可扩展:新增产品类别、新接入渠道、新增质检模型,均无需架构级改造;",
        "安全合规优先:国密算法、全量审计、监管直连是底线,任何优化不能以牺牲合规为代价。",
    ])

    add_heading(doc, "3.2 整体架构图", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "01_overall_architecture.png"),
               "图 3-1  线上线下双录一体化平台整体架构 (5 层 + 横切能力)")
    add_paragraph_text(doc,
        "整体架构采用经典的分层设计,自上而下分为 5 个层次,并配备 5 个横切能力中心,确保核心业务逻辑与平台能力解耦。")

    add_heading(doc, "3.2.1 5 个核心层", level=3)
    layers = [
        ["接入层",       "线上 H5/小程序、线下一体机、PAD 移动展业、网点 PC 端"],
        ["流程编排层",   "BPMN 流程引擎、状态机、异常分支处理、回退补偿机制"],
        ["能力中台层",   "统一话术中心、双录引擎、智能质检引擎、电子签、风评引擎"],
        ["数据治理层",   "统一订单中心、分布式事务协调器、客户主数据、影像存证、审计追溯"],
        ["基础底座层",   "国密加密、对象存储 OSS、消息中间件、统一身份/权限、监管上报通道"],
    ]
    add_table_styled(doc, ["层次", "包含能力"], layers, col_widths=[1.5, 5.5])

    add_heading(doc, "3.2.2 5 个横切能力", level=3)
    cross_caps = [
        ["统一话术中心", "话术/风险点/产品参数/合规要点统一管理,版本化 + MD5 校验"],
        ["智能质检引擎", "规则 + ASR/NLP/情绪 多模态 AI 质检,结果实时回传"],
        ["全链路追踪",   "预约→双录→签单→归档全链路埋点,任意节点可定位"],
        ["合规审计",     "监管接口直连 · 全量留痕 · 可回放,可司法采信"],
        ["可视化运营",   "管理者驾驶舱 · 实时大屏 · 异常告警,管理决策有据可依"],
    ]
    add_table_styled(doc, ["横切能力", "说明"], cross_caps, col_widths=[1.5, 5.5])

    add_heading(doc, "3.3 技术选型", level=2)
    tech_stack = [
        ["前端",      "Vue 3 + TypeScript + WebRTC SDK;PAD Android/iOS 原生 SDK"],
        ["后端",      "Spring Boot 3 + Spring Cloud Alibaba + Nacos + Sentinel"],
        ["流程引擎",  "Camunda 7 (BPMN 2.0) + 自研状态机引擎"],
        ["数据库",    "MySQL 8.0 (MGR 集群) + Redis 6.2 Cluster + Elasticsearch 8.x"],
        ["消息",      "Apache Kafka 3.x (3 broker,异步事件)"],
        ["存储",      "阿里云 OSS (3 副本 EC) + 冷归档 + 区块链 (Hyperledger Fabric)"],
        ["AI 能力",   "阿里云 ASR + 自研 NLP 模型 + 第三方情感分析 API"],
        ["加密",      "国密 SM2/SM3/SM4 芯片级加密 + 国产密码机"],
        ["监控",      "Prometheus + Grafana + 阿里云 ARMS + ELK 日志"],
    ]
    add_table_styled(doc, ["技术域", "选型/方案"], tech_stack, col_widths=[1.5, 5.5])

    add_heading(doc, "3.4 系统部署架构", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "09_deployment.png"),
               "图 3-2  系统部署架构 (3 个逻辑区 + 9 个核心服务)")
    add_paragraph_text(doc,
        "部署采用 DMZ/生产/办公 三网隔离架构,所有通信走国密加密 + mTLS,数据库采用 MGR 多主 + Redis 6 节点分片 + 区块链 4 节点共识,确保单点故障不影响业务连续性。")

    add_page_break(doc)

    # ══════ Chapter 4: 接入层详细设计 ══════
    add_heading(doc, "第 4 章  接入层详细设计", level=1)
    add_paragraph_text(doc,
        "接入层承担所有客户与客户经理的交互入口,需在统一业务流程规范下适配不同终端的硬件能力、网络条件与操作习惯。")

    add_heading(doc, "4.1 线上 H5/小程序", level=2)
    add_bullet_list(doc, [
        "WebRTC + 自研音视频 SDK,支持实时音视频采集、双向对讲、本地预览;",
        "可信时间戳嵌入视频帧水印,防止录屏伪造;",
        "断点续传机制:网络中断后自动重连,断点处继续录制,避免重做;",
        "双流录制:人脸画面 + 证件画面分轨保存,便于回溯;",
        "轻量化部署,首次接入 < 3MB,启动 < 2 秒,适配中低端机型。",
    ])

    add_heading(doc, "4.2 线下一体机", level=2)
    add_bullet_list(doc, [
        "嵌入式定制设备,防拆卸、防篡改,仅运行受信任的应用程序;",
        "国密 SM4 加密芯片,密钥物理隔离,无法软件提取;",
        "本地暂存 + 4G/有线双通道实时回传,断网不丢数据;",
        "高清摄像头 + 证件扫描仪,支持身份证、银行卡、合同 A4 全自动识别;",
        "活体检测 + OCR 联动,核身 5 秒内完成。",
    ])

    add_heading(doc, "4.3 PAD 移动展业", level=2)
    add_bullet_list(doc, [
        "iOS/Android 原生 SDK,深度集成系统权限;",
        "外拓场景优化:室外强光、网络弱、临时场地,均有专门适配;",
        "GPS + WiFi 定位,佐证业务发生地点,防止异地虚假办理;",
        "完成后立即加密上传,本地不留任何明文客户资料;",
        "支持客户经理远程接单、跨网点支援,业务灵活度大幅提升。",
    ])

    add_heading(doc, "4.4 网点 PC 端", level=2)
    add_bullet_list(doc, [
        "营业网点大屏终端,面向 60+ 客户群体的简化版操作;",
        "高对比度 UI + 大字体 + 语音辅助,适老化友好;",
        "紧急情况下可由大堂经理协助操作,但所有关键节点必须客户本人确认。",
    ])

    add_page_break(doc)

    # ══════ Chapter 5: 流程编排层 ══════
    add_heading(doc, "第 5 章  流程编排层", level=1)
    add_paragraph_text(doc,
        "流程编排层是整个平台的中枢,基于 BPMN 2.0 标准 + 自研状态机引擎实现。所有业务节点、跳转条件、异常分支、回退补偿均通过可视化流程定义,业务方可低代码维护。")

    add_heading(doc, "5.1 流程引擎", level=2)
    add_bullet_list(doc, [
        "基于 Camunda 7 Modeler 二次开发,支持复杂网关、排他网关、并行网关、子流程;",
        "支持动态参数注入,话术版本、产品参数、客户等级均可作为流程变量;",
        "支持多租户隔离,不同业务线、不同分行可独立配置流程;",
        "流程定义版本化管理,任意历史版本可一键回滚。",
    ])

    add_heading(doc, "5.2 状态机", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "03_state_machine.png"),
               "图 5-1  双录订单状态机 (S0→S6)")
    add_paragraph_text(doc,
        "订单状态机定义 7 个核心状态(S0-S6),状态流转严格受控,任一非法跳转均被拒绝。状态变更通过 Saga 编排器持久化,支持任意时刻回放与回滚。")

    add_heading(doc, "5.3 Saga 分布式事务", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "04_saga_transaction.png"),
               "图 5-2  Saga 分布式事务时序")
    add_paragraph_text(doc,
        "双录流程涉及 6 个核心事务步骤(T1-T6),采用 Saga 模式实现分布式事务一致性。每一步都有对应的补偿动作,任一失败即触发自动回滚,确保不会出现半成品订单。")

    add_heading(doc, "5.4 关键检查点(Gate)", level=2)
    add_table_styled(doc,
        ["Gate", "名称", "不通过后果"],
        [
            ["G1", "实名核身通过",   "阻断后续流程,提示客户重新核身"],
            ["G2", "风评等级与产品匹配", "禁止购买不匹配产品,建议降级或加风险揭示"],
            ["G3", "话术 100% 执行",   "流程挂起,客户经理需重新执行遗漏节点"],
            ["G4", "客户意愿明确确认", "无法签署,需重新确认或终止办理"],
            ["G5", "质检分数 ≥ 70",   "自动转人工复检,不可直接归档"],
        ],
        col_widths=[0.8, 2.2, 4.0])

    add_page_break(doc)

    # ══════ Chapter 6: 能力中台层 ══════
    add_heading(doc, "第 6 章  能力中台层", level=1)
    add_paragraph_text(doc,
        "能力中台层是平台的核心,提供 5 大能力中心(话术/双录/质检/电子签/风评),所有业务能力以微服务形式注册到服务中心,统一对外提供 API。")

    add_heading(doc, "6.1 统一话术中心", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "05_script_atomic.png"),
               "图 6-1  话术原子化结构")

    add_heading(doc, "6.1.1 原子化拆解", level=3)
    add_paragraph_text(doc,
        "话术被原子化拆解为 N1~Nn 共 6 类节点(问候、产品告知、风险揭示、适当性匹配、犹豫期、客户意愿)。每类节点独立配置,产品方更新话术不影响其他节点,降低维护成本。")

    add_heading(doc, "6.1.2 强制约束", level=3)
    add_bullet_list(doc, [
        "客户端必须完整听读完整陈述,UI 强制进度条 + 倒计时;",
        "强制确认位必须由客户本人口头表达,ASR 校验包含「是/确认/同意」等关键词;",
        "节点跳过即阻断,事后 100% 回放校验,人工不可干预。",
    ])

    add_heading(doc, "6.1.3 版本管理", level=3)
    add_bullet_list(doc, [
        "按「产品 + 监管版本 + 生效日期」三维管理;",
        "支持灰度发布:先向 5% 流量推送,72 小时无异常后全量;",
        "旧版本可查询、可回放,不可被新版本覆盖删除;",
        "所有版本变更产生审计日志,可追溯至具体操作人。",
    ])

    add_heading(doc, "6.1.4 统一推送", level=3)
    add_paragraph_text(doc,
        "线上 SDK、线下一体机、PAD 端均通过统一 API 拉取话术,服务端用 MD5 + 版本号确保完整性,客户端本地缓存,断网情况下也能继续使用。")

    add_heading(doc, "6.2 双录引擎", level=2)
    add_bullet_list(doc, [
        "音视频采集:WebRTC 协议 + 自研 H.264/H.265 编码器,带宽自适应 1-4Mbps;",
        "可信时间戳:嵌入视频帧水印,国家授时中心 NTP 同步,精度 ±0.5s;",
        "本地暂存 + 实时回传,断网时本地保留 24 小时;",
        "录制完成后自动生成 SHA-256 指纹,作为后续校验基准;",
        "GPU 加速实时水印叠加,不显著增加客户端 CPU 负担。",
    ])

    add_heading(doc, "6.3 智能质检引擎", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "07_quality_pipeline.png"),
               "图 6-2  智能质检三层流水线")

    add_heading(doc, "6.3.1 L1 规则层", level=3)
    add_paragraph_text(doc,
        "100+ 规则模板覆盖:必读项检测、必答项校验、必签字校验、风险关键词、敏感词过滤。规则引擎基于 Drools,执行时间 < 0.5s,准确率 > 99%。")

    add_heading(doc, "6.3.2 L2 AI 智能层", level=3)
    add_bullet_list(doc, [
        "ASR 转写:阿里云语音识别,中文准确率 > 98%,方言自适应;",
        "NLP 意图识别:基于自研 BERT 模型的金融领域意图分类器;",
        "情感分析:检测客户犹豫、紧张、强制话术诱导等情绪;",
        "图像识别:人脸活体、证件真伪、签名笔迹;",
        "声纹识别:防止录制过程换人;",
        "多模态融合:30s 内完成综合评分,准确率 > 95%。",
    ])

    add_heading(doc, "6.3.3 L3 人工复核层", level=3)
    add_paragraph_text(doc,
        "高风险单 100% 复检,中风险抽样 30%,低风险 5% 抽检。争议件优先处理,所有结果可申诉,申诉通道 24 小时内响应。")

    add_heading(doc, "6.4 电子签", level=2)
    add_bullet_list(doc, [
        "采用合规电子签名服务,符合《电子签名法》要求;",
        "客户人脸核身 + 短信验证码 + 意愿确认三因子;",
        "CA 数字证书 + 时间戳服务,签名具备司法效力;",
        "电子合同哈希值上链区块链,事后不可篡改;",
        "支持手写板原笔迹签名,法律效力等同纸质签名。",
    ])

    add_heading(doc, "6.5 风险评估引擎", level=2)
    add_bullet_list(doc, [
        "标准化 KYC 问卷,涵盖基本信息、投资经验、风险偏好、资金性质;",
        "智能匹配算法:根据回答自动评定 C1-C5 五个风险等级;",
        "风险等级与产品等级严格匹配,等级不匹配时强制升级风险揭示;",
        "评估结果有效期 12 个月,过期必须重新评估;",
        "评估过程与购买流程必须在同一次双录中,跨次评估视作失效(监管红线)。",
    ])

    add_page_break(doc)

    # ══════ Chapter 7: 数据治理层 ══════
    add_heading(doc, "第 7 章  数据治理层", level=1)
    add_paragraph_text(doc,
        "数据治理层是平台的「记忆系统」,统一管理订单、客户、话术、风评、质检、合同等核心数据,确保数据一致性、可追溯、合规存档。")

    add_heading(doc, "7.1 核心数据模型", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "10_data_model.png"),
               "图 7-1  核心数据模型 ER 图")

    add_heading(doc, "7.1.1 实体清单", level=3)
    entities = [
        ["客户 (Customer)",     "全行统一客户主数据,跨业务线打通,主键 customer_id"],
        ["订单 (Order)",         "双录业务订单,主键 order_id,关联客户、产品、风评、合同"],
        ["双录会话 (Session)",   "一笔订单对应一个或多个双录会话,主键 session_id"],
        ["话术 (Script)",        "原子化话术模板,主键 script_id,JSON 存储节点内容"],
        ["风评 (RiskAssess)",   "KYC 评估结果,主键 assess_id,关联客户与订单"],
        ["质检 (Quality)",       "智能质检结果,主键 qa_id,关联会话与评分卡"],
        ["电子合同 (Contract)", "电子合同文件,主键 contract_id,签名证书 + 哈希"],
    ]
    add_table_styled(doc, ["实体", "说明"], entities, col_widths=[2.0, 5.0])

    add_heading(doc, "7.2 数据库设计要点", level=2)
    add_bullet_list(doc, [
        "主库 MySQL 8.0 采用 MGR 多主架构,3 节点保证强一致性和高可用;",
        "分库分表策略:订单按创建时间月分库,客户按 customer_id hash 分 64 库;",
        "索引设计:所有外键 + 业务查询字段均建索引,慢 SQL < 0.1%;",
        "软删除:所有核心表保留 deleted_at,数据可恢复,可审计;",
        "WORM 存储:合规相关表启用 Write Once Read Many,180 天不可修改。",
    ])

    add_heading(doc, "7.3 分布式事务实现", level=2)
    add_paragraph_text(doc,
        "采用 Saga 模式而非 2PC,原因:1)双录流程长事务(分钟级),2PC 锁资源代价过高;2)涉及外部服务较多,2PC 协调复杂;3)Saga 补偿机制天然契合「任一失败回滚」的业务诉求。")

    add_heading(doc, "7.4 数据一致性保障", level=2)
    add_bullet_list(doc, [
        "强一致:订单主表、订单状态、订单日志通过同库事务保证;",
        "最终一致:风评、质检、合同通过事件总线异步同步,30s 内完成;",
        "补偿机制:每一步都对应反向操作,失败时按倒序回滚;",
        "幂等设计:所有外部调用支持重试,通过唯一业务键防重;",
        "数据校验:T+1 自动核对所有异构存储,异常时触发告警。",
    ])

    add_page_break(doc)

    # ══════ Chapter 8: 业务流程详细设计 ══════
    add_heading(doc, "第 8 章  业务流程详细设计", level=1)

    add_heading(doc, "8.1 端到端主流程", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "02_end_to_end_flow.png"),
               "图 8-1  端到端业务流程 (预约 → 归档)")
    add_paragraph_text(doc,
        "双录业务主流程分为 8 个阶段,每个阶段对应明确的责任主体、产出物和质量门禁。")

    stages_detail = [
        ["01 预约",  "客户身份预校验 + 产品匹配", "客户/客户经理", "预约单"],
        ["02 排程",  "客户经理排期 + 渠道分配", "客户经理", "排程单"],
        ["03 核身",  "证件 OCR + 活体检测 + 双签", "客户", "核身结果"],
        ["04 风评",  "KYC 问卷 + 风险等级评定", "客户", "风评报告"],
        ["05 双录",  "话术执行 + 音视频同步录制", "客户/客户经理", "双录视频"],
        ["06 签约",  "电子合同 + CA 数字签名", "客户", "电子合同"],
        ["07 质检",  "AI 智能 + 人工复核", "系统/人工", "质检报告"],
        ["08 归档",  "区块链存证 + 监管上报", "系统", "归档凭证"],
    ]
    add_table_styled(doc, ["阶段", "内容", "责任方", "产出物"],
                     stages_detail, col_widths=[1.0, 2.5, 1.5, 2.0])

    add_heading(doc, "8.2 异常分支流程", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "08_exception_handling.png"),
               "图 8-2  异常处理分类与恢复")

    add_heading(doc, "8.2.1 三类异常处理", level=3)
    exception_table = [
        ["技术异常", "网络断线/设备故障/音视频异常/OSS 上传失败",
         "本地缓存 + 断点续传 + 自动告警,5s 内自动恢复"],
        ["流程异常", "客户拒答/话术中断/中途离席/二次确认失败",
         "重新进入该节点,客户经理介入,30s 内可继续"],
        ["合规异常", "风险不匹配/证件失效/非本人办理/高风险反洗钱",
         "阻断流程 + 升级审核 + 留痕上报,必须人工处理"],
    ]
    add_table_styled(doc, ["类型", "场景示例", "处理策略"],
                     exception_table, col_widths=[1.2, 3.0, 2.8])

    add_heading(doc, "8.2.2 客户保护机制", level=3)
    add_paragraph_text(doc,
        "任何异常都不能让客户「卡在流程里」。6 重保护机制确保客户体验:")
    add_bullet_list(doc, [
        "多端续接:线上中断可在线下继续,反之亦然;",
        "暂存恢复:本地加密暂存 24 小时,网络恢复自动续传;",
        "远程协助:客户经理远程接管 + 屏幕共享;",
        "二次预约:本次失败自动 7 天内可重新预约,无需人工;",
        "人工兜底:复杂情况转 955xx 人工坐席;",
        "短信进度:每完成 1 个节点,短信通知客户进度。",
    ])

    add_heading(doc, "8.3 渠道分支策略", level=2)
    add_paragraph_text(doc,
        "同一业务流程在不同渠道的载体不同,根据场景自动适配:")
    add_table_styled(doc,
        ["渠道", "时长", "客户角色", "适用场景"],
        [
            ["线上 H5/小程序", "5-8 分钟",  "客户自助",      "简单产品、年轻客户、远程办理"],
            ["线下一体机",     "15-20 分钟","客户经理陪同", "高净值客户、复杂产品、需面签"],
            ["PAD 移动展业",   "10-15 分钟","客户经理上门", "外拓场景、社区营销、企业团办"],
        ],
        col_widths=[1.5, 1.2, 1.8, 2.5])

    add_page_break(doc)

    # ══════ Chapter 9: 话术模板详细设计 ══════
    add_heading(doc, "第 9 章  话术模板详细设计", level=1)
    add_paragraph_text(doc,
        "话术模板是双录合规的核心载体,本平台针对 4 类产品(保险/理财/基金/风评)分别设计标准化话术,所有话术经过法务、合规、业务三方评审。")

    add_heading(doc, "9.1 保险产品话术", level=2)
    insurance_script = [
        ["N1 问候核身", "您好,我是 XX 银行客户经理 XX,本次双录将全程录音录像,作为合规凭证保存 10 年,请问您同意吗?"],
        ["N2 产品告知", "本产品为终身寿险(分红型),保险责任以合同条款为准,您是否清楚?"],
        ["N3 风险揭示", "分红型保险收益不确定,演示利率非保证,可能为零,您是否充分理解?"],
        ["N4 犹豫期",   "您有 15 天犹豫期,期内退保仅扣 10 元工本费,是否知晓?"],
        ["N5 健康告知", "请如实告知健康状况,隐瞒可能影响理赔,是否如实?"],
        ["N6 签字确认", "以上内容均为您的真实意愿表达,是否确认投保?"],
    ]
    add_table_styled(doc, ["节点", "标准话术"], insurance_script, col_widths=[1.5, 5.5])

    add_heading(doc, "9.2 理财产品话术", level=2)
    wealth_script = [
        ["N1 风险匹配", "经评估您的风险等级为 C3(稳健型),本产品风险等级 C2,匹配通过,是否知晓?"],
        ["N2 产品要素", "本产品为封闭式净值型理财,期限 365 天,业绩比较基准 3.5%-4.2%。"],
        ["N3 收益说明", "业绩比较基准非保证收益,实际可能低于基准,本金不保证,您是否知晓?"],
        ["N4 资金用途", "本产品募集资金主要投资于债券和同业存单。"],
        ["N5 流动性",   "产品封闭期内不可赎回,是否影响您的资金安排?"],
        ["N6 确认",     "请确认知悉风险与产品要素,是否继续购买?"],
    ]
    add_table_styled(doc, ["节点", "标准话术"], wealth_script, col_widths=[1.5, 5.5])

    add_heading(doc, "9.3 基金产品话术", level=2)
    fund_script = [
        ["N1 适当性",   "经评估您的风险等级为 C4(成长型),本基金风险等级 C4,匹配通过,是否知晓?"],
        ["N2 基金类型", "本基金为股票型开放式基金,股票仓位不低于 80%。"],
        ["N3 净值波动", "基金净值随市场波动,可能出现本金损失,极端情况下亏损超过 30%,您是否理解?"],
        ["N4 费率说明", "申购费 1.5%,管理费 1.5%/年,托管费 0.25%/年,持有 7 天内赎回 1.5% 罚息。"],
        ["N5 历史业绩", "过往业绩不预示未来表现,投资需谨慎。"],
        ["N6 确认",     "请确认已阅读《招募说明书》《风险揭示书》,是否确认申购?"],
    ]
    add_table_styled(doc, ["节点", "标准话术"], fund_script, col_widths=[1.5, 5.5])

    add_heading(doc, "9.4 风险评估话术", level=2)
    risk_script = [
        ["Q1 知情同意", "为给您匹配合适的产品,需先做风险评估,回答将作为销售依据,您是否同意?"],
        ["Q2 收入询问", "请告知您家庭的年收入范围:A.10 万以下 B.10-30 万 C.30-100 万 D.100 万以上"],
        ["Q3 投资经验", "您过往投资过哪些类型?股票/基金/银行理财/不动产/无,投资年限?"],
        ["Q4 风险态度", "若投资亏损 20%,您会:A.无法接受 B.焦虑但持有 C.可接受 D.加仓抄底"],
        ["Q5 流动性",   "这笔资金的预计持有期限?短期(1 年内)/中期(1-3 年)/长期(3 年以上)"],
        ["Q6 风险揭示", "评估结果:您的风险等级为 C3(稳健型),适合 R3 及以下风险产品,是否知晓?"],
        ["Q7 二次确认", "以上评估基于您的真实意愿,如有虚假需承担相应责任,是否确认?"],
    ]
    add_table_styled(doc, ["节点", "标准话术"], risk_script, col_widths=[1.5, 5.5])

    add_callout(doc, "细节点", "风评双录关联:评估结果与产品购买必须在同一次双录中,跨次评估视作失效(监管红线)。")

    add_page_break(doc)

    # ══════ Chapter 10: 视频合规方案 ══════
    add_heading(doc, "第 10 章  视频合规方案", level=1)
    add_paragraph_text(doc,
        "视频合规是双录业务的「生命线」,本方案从录制、加密、存储、存证、监管对接全链路确保视频证据的法律效力。")

    add_heading(doc, "10.1 录制规范", level=2)
    add_bullet_list(doc, [
        "音视频规格:H.264/H.265 编码,720P/1080P 分辨率,音频 16kHz 采样;",
        "时间戳:视频帧嵌入可信时间戳,精度 ±0.5s,以国家授时中心 NTP 同步;",
        "人脸画面:客户 + 客户经理双画面 PIP(画中画),清晰度可分辨五官;",
        "证件画面:同步展示客户身份证、银行卡、合同,确保与视频同步;",
        "全程录制:从问候开始到签字确认结束,中间不可中断,所有节点必须连贯。",
    ])

    add_heading(doc, "10.2 视频加密与存储", level=2)
    add_figure(doc, os.path.join(DIAG_DIR, "06_compliance_chain.png"),
               "图 10-1  视频合规证据链")
    add_paragraph_text(doc,
        "视频采用国密 SM4 芯片级加密 → 对象存储 → SHA-256 指纹 → 区块链存证 → 监管上报全链路处理。")
    add_bullet_list(doc, [
        "加密:录制完成后立即 SM4 加密,密钥与视频分离管理;",
        "存储:OSS 3 副本 EC,按热/温/冷分层,合规要求 10 年留存;",
        "指纹:SHA-256 哈希校验,任何篡改都会导致指纹不匹配;",
        "存证:哈希值 + 关键元数据上链 Hyperledger Fabric,司法可采信;",
        "监管:与银保监/证监会监管平台直连,T+1 自动报送。",
    ])

    add_heading(doc, "10.3 存储分层策略", level=2)
    add_table_styled(doc,
        ["层级", "保留期限", "存储类型", "访问频率", "成本"],
        [
            ["热存储", "0-90 天",   "OSS 标准型",  "高频访问", "高"],
            ["温存储", "90 天-3 年","OSS 低频型",  "偶发查询", "中"],
            ["冷存储", "3-10 年",   "OSS 归档型",  "极少查询", "低"],
        ],
        col_widths=[1.0, 1.5, 1.5, 1.5, 1.5])

    add_heading(doc, "10.4 调阅与审计", level=2)
    add_bullet_list(doc, [
        "调阅权限:仅合规、稽核、监管人员可调阅,且需双因素认证 + 审批;",
        "调阅记录:所有调阅行为留痕,操作人/时间/原因/范围全记录;",
        "司法取证:支持与司法鉴定机构对接,提供完整证据包;",
        "客户查询:客户本人可通过手机银行查询本人的双录视频(脱敏版);",
        "监管对接:监管检查时 30 分钟内可调取任何一笔订单的全套证据。",
    ])

    add_page_break(doc)

    # ══════ Chapter 11: 智能质检方案 ══════
    add_heading(doc, "第 11 章  智能质检方案", level=1)

    add_heading(doc, "11.1 质检规则库", level=2)
    add_paragraph_text(doc,
        "质检规则库是整个质检系统的基础,目前已沉淀 100+ 规则模板,按业务类型分类管理:")
    rule_categories = [
        ["通用规则", "必读项、必答项、必签字、风险关键词、敏感词"],
        ["保险规则", "犹豫期告知、如实告知、退保权利、责任免除"],
        ["理财规则", "业绩比较基准、封闭期、提前赎回罚息、资金用途"],
        ["基金规则", "风险等级匹配、净值波动警示、费率透明、风险揭示书"],
        ["流程规则", "核身通过、风评完成、签字完整、时长合理"],
    ]
    add_table_styled(doc, ["类别", "规则示例"], rule_categories, col_widths=[1.5, 5.5])

    add_heading(doc, "11.2 统一评分卡", level=2)
    add_paragraph_text(doc,
        "线上/线下使用同一张评分卡,五大维度综合打分,确保结果可比、可对账:")
    add_table_styled(doc,
        ["维度", "权重", "评分要点"],
        [
            ["话术完整度",   "30%", "是否覆盖所有必读项,是否漏读关键话术"],
            ["风险揭示准确度","25%", "风险点是否清晰、准确,客户是否理解"],
            ["客户确认清晰度","20%", "客户确认是否明确,有无含糊其辞"],
            ["音视频合规度", "15%", "画面/声音是否清晰,关键节点是否完整"],
            ["流程完整度",   "10%", "所有 Gate 是否通过,流程是否规范"],
        ],
        col_widths=[2.0, 1.0, 4.0])

    add_heading(doc, "11.3 质检结果分级处理", level=2)
    add_table_styled(doc,
        ["评分", "等级", "处理方式"],
        [
            ["≥ 90", "高分", "自动归档,无需复检"],
            ["70-89","中分", "AI 标记异常点 + 抽检 30%"],
            ["< 70", "低分", "100% 人工复检,不可直接归档"],
        ],
        col_widths=[1.5, 1.5, 4.0])

    add_heading(doc, "11.4 AI 模型管理", level=2)
    add_bullet_list(doc, [
        "模型版本化:每次训练产生新版本,AB 测试 7 天后全量;",
        "模型监控:实时监控准确率/召回率,异常自动告警;",
        "模型回滚:任一指标下降超过 5%,自动回滚到上一稳定版本;",
        "持续学习:人工复检结果回流到训练集,模型持续迭代;",
        "领域适配:针对不同业务(保险/理财/基金)训练专项模型,准确率提升 10%。",
    ])

    add_page_break(doc)

    # ══════ Chapter 12: 安全与合规 ══════
    add_heading(doc, "第 12 章  安全与合规", level=1)

    add_heading(doc, "12.1 数据安全", level=2)
    add_bullet_list(doc, [
        "传输加密:所有通信采用国密 SM2 + TLS 1.3 双重加密;",
        "存储加密:客户敏感信息(身份证/银行卡)字段级加密,密钥由硬件密码机管理;",
        "脱敏显示:非必要场景下身份证/手机号/银行卡脱敏显示;",
        "数据隔离:按业务线/分行/客户经理多维隔离,权限最小化;",
        "WORM 存储:合规关键数据 180 天内不可修改、不可删除。",
    ])

    add_heading(doc, "12.2 访问控制", level=2)
    add_bullet_list(doc, [
        "统一身份:基于全行 IAM 体系,单点登录 + 多因素认证;",
        "细粒度权限:RBAC 模型,权限可控制到字段级和操作级;",
        "API 网关:统一入口,限流、熔断、黑白名单;",
        "操作审计:所有操作产生审计日志,180 天留存,可追溯;",
        "特权账号:堡垒机 + 双人复核,杜绝单人操作风险。",
    ])

    add_heading(doc, "12.3 监管合规要求", level=2)
    add_bullet_list(doc, [
        "双录视频保存期限 ≥ 10 年(保险)/ 5 年(理财基金);",
        "电子签名符合《电子签名法》要求,具备司法效力;",
        "客户敏感信息保护符合《个人信息保护法》要求;",
        "反洗钱:大额/可疑交易实时上报,不留时间窗口;",
        "监管报送:T+1 自动报送,支持监管现场检查。",
    ])

    add_heading(doc, "12.4 应急与灾备", level=2)
    add_bullet_list(doc, [
        "RTO(恢复时间目标)≤ 30 分钟,RPO(数据丢失容忍)≤ 5 分钟;",
        "同城主备 + 异地灾备,数据库 MGR 多活;",
        "季度灾备演练,年度跨地域切换演练;",
        "应急响应预案:覆盖网络/数据库/应用/第三方服务 4 类故障;",
        "7×24 小时监控值守,异常 5 分钟内响应。",
    ])

    add_page_break(doc)

    # ══════ Chapter 13: 接口设计 ══════
    add_heading(doc, "第 13 章  接口设计", level=1)

    add_heading(doc, "13.1 接口设计原则", level=2)
    add_bullet_list(doc, [
        "RESTful 风格:URL 表示资源,HTTP 方法表示操作;",
        "统一响应格式:所有接口返回 {code, message, data, traceId};",
        "统一鉴权:OAuth 2.0 + JWT,所有接口必须 token 鉴权;",
        "幂等设计:写接口支持幂等键,避免重复操作;",
        "版本管理:URL 中带版本号 /v1/,平滑升级;",
        "限流熔断:接入 Sentinel,保护后端服务;",
        "全链路追踪:每个请求带 traceId,ELK 关联分析。",
    ])

    add_heading(doc, "13.2 核心 API 列表", level=2)
    api_list = [
        ["POST", "/api/v1/order/create",       "创建订单"],
        ["POST", "/api/v1/order/{id}/start",   "开始双录"],
        ["POST", "/api/v1/script/pull",        "拉取话术"],
        ["POST", "/api/v1/script/submit",      "提交话术执行结果"],
        ["POST", "/api/v1/session/upload",     "上传双录视频分片"],
        ["POST", "/api/v1/session/finalize",   "完成双录会话"],
        ["POST", "/api/v1/sign/apply",         "申请电子签名"],
        ["POST", "/api/v1/sign/verify",        "校验电子签名"],
        ["POST", "/api/v1/quality/audit",      "触发质检"],
        ["GET",  "/api/v1/quality/{id}/result","查询质检结果"],
        ["POST", "/api/v1/assess/submit",      "提交风评答案"],
        ["GET",  "/api/v1/order/{id}/status",  "查询订单状态"],
    ]
    add_table_styled(doc, ["方法", "URL", "说明"], api_list,
                     col_widths=[0.8, 3.0, 3.2])

    add_heading(doc, "13.3 异步消息协议", level=2)
    add_paragraph_text(doc,
        "双录业务存在大量异步场景(质检、监管上报、消息通知),采用 Kafka 消息队列实现。")
    add_bullet_list(doc, [
        "Topic 设计:dr.qual.completed(质检完成)、dr.sign.done(签约完成)、dr.archive.done(归档完成);",
        "消息格式:CloudEvents 1.0 规范,包含事件源、时间、类型、数据;",
        "消费模式:至少一次投递 + 消费方幂等;",
        "失败重试:指数退避,最大重试 3 次,失败入死信队列;",
        "监控:消费延迟、消费成功率,异常实时告警。",
    ])

    add_page_break(doc)

    # ══════ Chapter 14: 实施计划 ══════
    add_heading(doc, "第 14 章  实施计划", level=1)

    add_heading(doc, "14.1 项目阶段", level=2)
    add_table_styled(doc,
        ["阶段", "时间", "里程碑", "交付物"],
        [
            ["P0 基础",  "M1-M2",   "架构落地",     "基础平台、统一身份、对象存储"],
            ["P1 话术",  "M3-M4",   "话术中心",     "话术原子化、版本管理、强制执行"],
            ["P2 双录",  "M5-M7",   "双录引擎",     "音视频采集、加密、上传、存证"],
            ["P3 质检",  "M8-M10",  "智能质检",     "三层质检流水线、统一评分卡"],
            ["P4 数据",  "M11-M12", "数据治理",     "分布式事务、状态机、数据一致性"],
            ["P5 上线",  "M13-M14", "试运行",       "灰度发布、全量上线、监控完善"],
        ],
        col_widths=[1.2, 1.0, 1.5, 3.3])

    add_heading(doc, "14.2 资源需求", level=2)
    add_bullet_list(doc, [
        "开发团队:产品经理 2 人,前端 2 人,后端 6 人,AI 2 人,测试 3 人,运维 1 人;",
        "硬件资源:应用服务器 20 台(8C16G),数据库服务器 6 台(32C128G),对象存储 100TB;",
        "第三方服务:ASR/NLP 接口调用额度,CA 证书年费,区块链节点 4 台;",
        "预算:首年约 1500 万(含人力/硬件/软件/培训),后续每年运维 300 万。",
    ])

    add_heading(doc, "14.3 风险评估", level=2)
    risk = [
        ["技术风险", "音视频 SDK 兼容性", "中", "选型阶段充分测试,准备降级方案"],
        ["技术风险", "AI 模型准确率",     "中", "初期保留人工兜底,持续优化"],
        ["合规风险", "监管政策变化",       "高", "建立监管跟踪机制,敏捷响应"],
        ["业务风险", "客户接受度",         "中", "前期用户研究 + UAT 充分验证"],
        ["项目风险", "跨部门协作",         "中", "成立 PMO,周例会 + 风险升级机制"],
        ["运维风险", "上线故障",           "中", "灰度发布 + 完善监控 + 应急演练"],
    ]
    add_table_styled(doc, ["类型", "风险项", "等级", "应对措施"],
                     risk, col_widths=[1.0, 2.0, 0.8, 3.2])

    add_page_break(doc)

    # ══════ Appendix A: Data Dictionary ══════
    add_heading(doc, "附录 A  数据字典", level=1)
    add_paragraph_text(doc, "本附录列出本平台核心数据表的字段定义,作为开发、测试、运维的共同参考。")

    add_heading(doc, "A.1 客户表 (t_customer)", level=2)
    add_table_styled(doc,
        ["字段", "类型", "必填", "说明"],
        [
            ["customer_id",   "BIGINT",     "是", "客户唯一标识,全局主键"],
            ["name",          "VARCHAR(64)","是", "客户姓名,加密存储"],
            ["id_type",       "TINYINT",    "是", "证件类型 1-身份证 2-护照 3-军官证"],
            ["id_no",         "VARCHAR(32)","是", "证件号,SM4 加密"],
            ["id_expire_date","DATE",       "否", "证件有效期,用于过期校验"],
            ["risk_level",    "CHAR(2)",    "是", "风险等级 C1-C5,默认 C1"],
            ["risk_expire_at","DATETIME",   "是", "风险评估过期时间,12 个月有效"],
            ["created_at",    "DATETIME",   "是", "创建时间"],
            ["updated_at",    "DATETIME",   "是", "最后修改时间"],
            ["deleted_at",    "DATETIME",   "否", "软删除时间"],
        ],
        col_widths=[1.5, 1.2, 0.7, 3.6])

    add_heading(doc, "A.2 订单表 (t_order)", level=2)
    add_table_styled(doc,
        ["字段", "类型", "必填", "说明"],
        [
            ["order_id",      "BIGINT",     "是", "订单唯一标识,全局主键"],
            ["customer_id",   "BIGINT",     "是", "客户 ID,外键"],
            ["product_id",    "BIGINT",     "是", "产品 ID,外键"],
            ["product_type",  "TINYINT",    "是", "产品类型 1-保险 2-理财 3-基金"],
            ["amount",        "DECIMAL(18,2)","是", "购买金额,分单位"],
            ["state",         "TINYINT",    "是", "订单状态 S0-S6,默认 S0"],
            ["channel",       "TINYINT",    "是", "渠道 1-H5 2-一体机 3-PAD 4-PC"],
            ["sales_user_id", "BIGINT",     "否", "客户经理 ID"],
            ["created_at",    "DATETIME",   "是", "创建时间"],
            ["completed_at",  "DATETIME",   "否", "完成时间"],
        ],
        col_widths=[1.5, 1.2, 0.7, 3.6])

    add_page_break(doc)

    # ══════ Appendix B: Script List ══════
    add_heading(doc, "附录 B  话术模板清单", level=1)
    add_paragraph_text(doc, "本附录列出当前已上线的全部话术模板,供业务方查阅。每个模板含版本号、生效日期、节点数。")
    add_table_styled(doc,
        ["模板编码", "产品", "版本", "生效日期", "节点数", "状态"],
        [
            ["INS-LIFE-V3",   "终身寿险",     "V3.2", "2026-01-15", "6", "已上线"],
            ["INS-HEALTH-V2", "健康险",       "V2.1", "2026-03-01", "6", "已上线"],
            ["WEALTH-CLOSED-V5","封闭式理财", "V5.0", "2026-05-20", "6", "已上线"],
            ["WEALTH-OPEN-V3","开放式理财",   "V3.4", "2026-04-10", "6", "已上线"],
            ["FUND-EQ-V4",    "股票型基金",   "V4.1", "2026-06-01", "6", "已上线"],
            ["FUND-BOND-V3",  "债券型基金",   "V3.0", "2026-02-15", "6", "已上线"],
            ["RISK-KYC-V6",   "KYC 风险评估", "V6.0", "2026-07-01", "7", "已上线"],
            ["INS-ANNUITY-V1","年金险",       "V1.0", "2026-07-15", "6", "灰度中"],
            ["WEALTH-MIXED-V1","混合型理财",  "V1.0", "2026-08-01", "6", "灰度中"],
        ],
        col_widths=[1.6, 1.5, 0.8, 1.2, 0.8, 1.1])

    add_page_break(doc)

    # ══════ Appendix C: Abbreviations ══════
    add_heading(doc, "附录 C  缩略语表", level=1)
    abbreviations = [
        ["API",   "Application Programming Interface",  "应用程序编程接口"],
        ["ASR",   "Automatic Speech Recognition",      "自动语音识别"],
        ["BPMN",  "Business Process Model and Notation","业务流程模型与标记法"],
        ["CA",    "Certificate Authority",             "数字证书认证机构"],
        ["CRM",   "Customer Relationship Management",  "客户关系管理"],
        ["DLT",   "Distributed Ledger Technology",      "分布式账本技术"],
        ["DMZ",   "Demilitarized Zone",                 "隔离区"],
        ["EC",    "Erasure Code",                       "纠删码"],
        ["H5",    "HTML 5 (移动端页面)",                "HTML 5"],
        ["IAM",   "Identity and Access Management",    "身份与访问管理"],
        ["KYC",   "Know Your Customer",                "了解你的客户"],
        ["MGR",   "MySQL Group Replication",           "MySQL 组复制"],
        ["NPS",   "Net Promoter Score",                 "净推荐值"],
        ["NLP",   "Natural Language Processing",        "自然语言处理"],
        ["OCR",   "Optical Character Recognition",     "光学字符识别"],
        ["OSS",   "Object Storage Service",             "对象存储服务"],
        ["REST",  "Representational State Transfer",    "表述性状态转移"],
        ["RPO",   "Recovery Point Objective",           "恢复点目标"],
        ["RTO",   "Recovery Time Objective",            "恢复时间目标"],
        ["SDK",   "Software Development Kit",           "软件开发工具包"],
        ["SLA",   "Service Level Agreement",            "服务等级协议"],
        ["SM",    "ShangMi (国密算法)",                 "中国国家密码算法"],
        ["TPS",   "Transactions Per Second",            "每秒事务数"],
        ["WAF",   "Web Application Firewall",           "Web 应用防火墙"],
        ["WORM",  "Write Once Read Many",               "一次写入多次读取"],
    ]
    add_table_styled(doc, ["缩写", "英文全称", "中文释义"],
                     abbreviations, col_widths=[1.0, 3.5, 2.5])

    # ══════ Save ══════
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Document saved: {OUT_PATH}")
    return OUT_PATH

if __name__ == "__main__":
    build()
